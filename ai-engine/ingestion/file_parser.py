"""
INGESTION MODULE — File Parser
Converts PDF, DOCX, and TXT files to clean raw text.
No AI involved here — pure file I/O.
"""

import io
import re
import chardet
from pathlib import Path
from loguru import logger

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not available. PDF parsing disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX parsing disabled.")


class FileParser:
    """
    Parses CV files into raw text strings.
    Supports: PDF, DOCX/DOC, TXT
    """

    SUPPORTED_MIME_TYPES = {
        "application/pdf": "_parse_pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "_parse_docx",
        "application/msword": "_parse_docx",
        "text/plain": "_parse_txt",
        "text/html": "_parse_txt",
    }

    def parse_bytes(self, content: bytes, mime_type: str, filename: str = "") -> str:
        """
        Parse file bytes into clean text.

        Args:
            content: Raw file bytes
            mime_type: MIME type string
            filename: Original filename (for extension fallback)

        Returns:
            Extracted raw text string

        Raises:
            ValueError: If file type unsupported or text too short
        """
        # Fallback: detect mime from extension
        if mime_type not in self.SUPPORTED_MIME_TYPES:
            ext = Path(filename).suffix.lower()
            mime_type = {
                ".pdf": "application/pdf",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".doc": "application/msword",
                ".txt": "text/plain",
            }.get(ext, mime_type)

        parser_method = self.SUPPORTED_MIME_TYPES.get(mime_type)
        if not parser_method:
            raise ValueError(f"Unsupported file type: {mime_type}")

        method = getattr(self, parser_method)
        raw_text = method(content)

        raw_text = self._clean_raw(raw_text)

        if len(raw_text.strip()) < 50:
            raise ValueError(
                "Extracted text is too short. "
                "The file may be image-based (scanned PDF) or empty."
            )

        logger.debug(f"Parsed {len(raw_text)} characters from {mime_type}")
        return raw_text

    def parse_file(self, file_path: str) -> str:
        """Parse a file from disk path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        with open(file_path, "rb") as f:
            content = f.read()

        ext = path.suffix.lower()
        mime_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
        }
        mime_type = mime_map.get(ext, "text/plain")
        return self.parse_bytes(content, mime_type, path.name)

    # ─── Parsers ──────────────────────────────────────────────────────────────

    def _parse_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber.
        Handles multi-column layouts, headers/footers, and tables.
        """
        if not PDF_AVAILABLE:
            raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Extract plain text
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text:
                        text_parts.append(page_text)

                    # Also extract table contents
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = " | ".join(
                                cell.strip() if cell else "" for cell in row
                            )
                            if row_text.strip():
                                text_parts.append(row_text)

                except Exception as e:
                    logger.warning(f"PDF page {page_num + 1} extraction error: {e}")
                    continue

        return "\n".join(text_parts)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX preserving paragraph structure."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

        doc = DocxDocument(io.BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))

        return "\n".join(parts)

    def _parse_txt(self, content: bytes) -> str:
        """Parse plain text, detecting encoding automatically."""
        detected = chardet.detect(content)
        encoding = detected.get("encoding") or "utf-8"
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            return content.decode("utf-8", errors="replace")

    # ─── Cleaning ─────────────────────────────────────────────────────────────

    def _clean_raw(self, text: str) -> str:
        """Basic text normalization — preserve structure for LLM."""
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Remove null bytes and control chars (keep tabs and newlines)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Collapse more than 3 consecutive newlines
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        # Collapse multiple spaces (but not newlines)
        text = re.sub(r"[ \t]{3,}", "  ", text)
        return text.strip()


# Module-level singleton
file_parser = FileParser()
